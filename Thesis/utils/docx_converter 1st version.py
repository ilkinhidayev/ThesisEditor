from docx import Document

def convert_to_guidelines(docx_path):
    # Upload Docx Files
    doc = Document(docx_path)

    # Converting and modifying file
    for paragraph in doc.paragraphs:
        # For instance, this adds 10 times space to the end of each paragraph
        paragraph.add_run("ilkin" * 10)

    # Save new file
    converted_path = docx_path.replace('.docx', '_converted.docx')
    doc.save(converted_path)

    return converted_path
